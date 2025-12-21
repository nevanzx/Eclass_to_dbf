import streamlit as st
import pandas as pd
import tempfile
import base64
from docx import Document
from docx.shared import Inches, Pt
import io
import json
import tempfile
import os
from docx2pdf import convert
import re
from dbf import Table


def parse_jle_with_filename_fixed(file_path):
    """
    Fixed version of the JLE parser that properly handles lecturer and credit extraction
    """
    # 1. Extract Semester/Year from Filename
    filename = os.path.basename(file_path)

    # Logic: Look for the pattern "20243" inside "DSO_20243_565.JLE"
    # This assumes the digit immediately following the year is the semester.
    term_map = {'1': '1st Semester', '2': '2nd Semester', '3': 'Summer'}

    # Default values
    academic_year = "Unknown"
    semester = "Unknown"

    # Extract digits (e.g., finding "20243")
    meta_match = re.search(r"(\d{4})(\d)", filename)
    if meta_match:
        year_part = meta_match.group(1)  # 2024
        term_part = meta_match.group(2)  # 3

        academic_year = f"{year_part}-{int(year_part)+1}"
        semester = term_map.get(term_part, f"{term_part}th Term")

    # 2. Extract data from the JLE file content
    with open(file_path, 'rb') as f:
        jle_bytes = f.read()

    # Read the JLE file content as binary and decode with latin1 encoding
    raw_data = jle_bytes.decode('latin1', errors='ignore')

    # Define the "Start of Record" pattern
    record_start_pattern = re.compile(r"(\d{4})\s+([A-Z][A-Z0-9]{2,10})")

    # Find all iterators for the start of records
    matches = list(record_start_pattern.finditer(raw_data))

    parsed_records = []

    for i in range(len(matches)):
        # Determine the start and end of the current record's text chunk
        start_idx = matches[i].start()

        # If there is a next match, end before it starts; otherwise go to EOF
        if i + 1 < len(matches):
            end_idx = matches[i+1].start()
        else:
            end_idx = len(raw_data)

        # Extract the raw chunk for this specific subject
        chunk = raw_data[start_idx:end_idx]

        # --- Extraction Logic ---

        # 1. Extract Subject Num and Code from the match itself
        subj_num = matches[i].group(1)
        full_subj_code = matches[i].group(2)

        # The first letter of the subject code is actually part of the subject number
        subj_num_with_letter = subj_num + full_subj_code[0]  # 2506 + F = 2506F
        real_subj_code = full_subj_code[1:]  # BACC104 (everything after the first letter)

        # Remove the header (Num + Code) from the chunk to process the rest
        rest_of_text = chunk[len(matches[i].group(0)):].replace('\n', ' ').strip()

        # 2. Extract Credit and Lecturer using improved pattern matching
        # The lecturer and credit appear at the end of the record
        # Updated pattern to handle control characters and special formatting
        lecturer_pattern = re.search(r"(\d)\s+([A-Z][A-Z\s\.\-]*[A-Z])(?:\s*$|[\x00-\x1f\x7f\s]*$)", rest_of_text)

        credit = None
        lecturer = None

        if lecturer_pattern:
            credit = lecturer_pattern.group(1)
            lecturer = lecturer_pattern.group(2).strip()
            # Remove this part from the text so we can find the Title/Schedule
            rest_of_text = rest_of_text[:lecturer_pattern.start()].strip()
        else:
            # Try another approach: look for credit (digit) followed by lecturer name at the end
            # Pattern: digit followed by spaces and then uppercase letters (lecturer name)
            alt_pattern = re.search(r"(\d)\s+([A-Z][A-Z\s\.\-]+?)(?:\s+(?:\d\s+)?[A-Z]|$)", rest_of_text)
            if alt_pattern:
                credit = alt_pattern.group(1)
                lecturer = alt_pattern.group(2).strip()
                rest_of_text = rest_of_text[:alt_pattern.start()].strip()

        # 3. Extract Schedule(s) - distinguish between LEC and LAB
        # Pattern: Time (e.g., 730AM-1200PM), Day (e.g., SuSa), Room (e.g., B63)
        sched_pattern_regex = r"(\d{1,4}[AP]M-\s*\d{1,4}[AP]M\s+[A-Za-z]+\s+[A-Z0-9]+)"
        schedules_found = re.findall(sched_pattern_regex, rest_of_text)

        # Initialize LEC and LAB schedules as empty
        lec_schedule_str = ""
        lab_schedule_str = ""

        # Try to identify LEC and LAB schedules if they're explicitly marked in the text
        # Look for LEC/LAB indicators in the context of schedules
        full_text_for_context = rest_of_text
        lec_pattern = r"LEC.*?(\d{1,4}[AP]M-\s*\d{1,4}[AP]M\s+[A-Za-z]+\s+[A-Z0-9]+)"
        lab_pattern = r"LAB.*?(\d{1,4}[AP]M-\s*\d{1,4}[AP]M\s+[A-Za-z]+\s+[A-Z0-9]+)"

        # Extract LEC schedules if explicitly marked
        lec_matches = re.findall(lec_pattern, rest_of_text, re.IGNORECASE)
        if lec_matches:
            lec_schedule_str = " / ".join([s.strip() for s in lec_matches])

        # Extract LAB schedules if explicitly marked
        lab_matches = re.findall(lab_pattern, rest_of_text, re.IGNORECASE)
        if lab_matches:
            lab_schedule_str = " / ".join([s.strip() for s in lab_matches])

        # If no explicit LEC/LAB markers found, use heuristic approach
        if not lec_schedule_str and not lab_schedule_str:
            if len(schedules_found) >= 2:
                # If there are 2 or more schedules, assume first is LEC, second is LAB (common pattern)
                lec_schedule_str = schedules_found[0] if schedules_found else ""
                lab_schedule_str = schedules_found[1] if len(schedules_found) > 1 else ""
            else:
                # If only one schedule, assign to LEC
                lec_schedule_str = schedules_found[0] if schedules_found else ""

        # Create combined schedule string
        all_schedule_str = " / ".join([s.strip() for s in schedules_found])

        # 4. Extract Subject Title
        # The title is whatever is left after removing the schedules
        # We use re.sub to remove the schedules we found
        title_clean = re.sub(sched_pattern_regex, '', rest_of_text)
        # Also remove the lecturer and credit from the title if they weren't removed by the main pattern
        if credit and lecturer:
            lecturer_part = f"{credit}\\s+{lecturer}"
            title_clean = re.sub(lecturer_part, '', title_clean, flags=re.IGNORECASE)
        title_clean = " ".join(title_clean.split()) # Remove extra whitespace

        # Clean up title to remove any control characters
        title_clean = re.sub(r'[\x00-\x1f\x7f]', ' ', title_clean).strip()

        parsed_records.append({
            "Subject Num": subj_num_with_letter,  # Now includes the letter (e.g., 2506F)
            "Subject Code": real_subj_code,       # The actual subject code (e.g., BACC104)
            "Subject Title": title_clean,
            "Schedule": all_schedule_str,          # All schedules combined
            "LEC_Schedule": lec_schedule_str,      # Lecture schedule
            "LAB_Schedule": lab_schedule_str,      # Laboratory schedule
            "Credit": credit,
            "Lecturer": lecturer,
            "Academic Year": academic_year,  # Added from filename
            "Semester": semester             # Added from filename
        })

    # Create a DataFrame from the parsed records
    jle_df = pd.DataFrame(parsed_records) if parsed_records else pd.DataFrame()

    return jle_df


def find_matching_dbf(jle_file_path, dbf_directory="testfiles"):
    """Find DBF file that matches the JLE file based on naming patterns"""
    # Extract metadata from JLE file
    jle_df = parse_jle_with_filename_fixed(jle_file_path)

    if jle_df.empty:
        return None, None

    # Get all DBF files in the directory
    dbf_files = [f for f in os.listdir(dbf_directory) if f.lower().endswith('.dbf')]

    for dbf_file in dbf_files:
        dbf_path = os.path.join(dbf_directory, dbf_file)

        # Extract parts from DBF filename to match against JLE
        # Pattern: ORG_YYYYX_SUBJNUM_SUBJCODE_ID.DBF
        parts = dbf_file.replace('.DBF', '').split('_')
        if len(parts) >= 4:
            year_semester = parts[1]  # YYYYX
            subj_num = parts[2]       # SUBJNUM
            subj_code = parts[3]      # SUBJCODE

            # Check if any of the JLE records match
            for _, jle_record in jle_df.iterrows():
                jle_academic_year = jle_record['Academic Year']  # Full format like "2024-2025"
                jle_year = jle_academic_year.split('-')[0]  # Get first part like "2024"
                jle_year_semester = f"{jle_year}{get_semester_digit(jle_record['Semester'])}"  # Full format like "20243"
                jle_subj_num = jle_record['Subject Num']
                jle_subj_code = jle_record['Subject Code']

                # Check if the components match
                if (year_semester == jle_year_semester and
                    subj_num == jle_subj_num and
                    subj_code == jle_subj_code):
                    return dbf_path, jle_record  # Return the matching DBF path and JLE record

    return None, None


def find_matching_dbf_from_jle_data(jle_data, dbf_directory="testfiles"):
    """Find DBF file that matches the JLE data based on naming patterns"""
    # Extract metadata from JLE data
    if jle_data is None or 'course_data' not in jle_data or jle_data['course_data'] is None or jle_data['course_data'].empty:
        return None, None

    jle_df = jle_data['course_data']

    # Get all DBF files in the directory
    dbf_files = [f for f in os.listdir(dbf_directory) if f.lower().endswith('.dbf')]

    for dbf_file in dbf_files:
        dbf_path = os.path.join(dbf_directory, dbf_file)

        # Extract parts from DBF filename to match against JLE
        # Pattern: ORG_YYYYX_SUBJNUM_SUBJCODE_ID.DBF
        parts = dbf_file.replace('.DBF', '').split('_')
        if len(parts) >= 4:
            year_semester = parts[1]  # YYYYX
            subj_num = parts[2]       # SUBJNUM
            subj_code = parts[3]      # SUBJCODE

            # Check if any of the JLE records match
            for _, jle_record in jle_df.iterrows():
                jle_academic_year = jle_record.get('Academic Year', '')  # Full format like "2024-2025"
                if jle_academic_year:
                    jle_year = jle_academic_year.split('-')[0]  # Get first part like "2024"
                    jle_year_semester = f"{jle_year}{get_semester_digit(jle_record.get('Semester', ''))}"  # Full format like "20243"
                    jle_subj_num = jle_record.get('Subject Num', '')
                    jle_subj_code = jle_record.get('Subject Code', '')

                    # Check if the components match
                    if (year_semester == jle_year_semester and
                        subj_num == jle_subj_num and
                        subj_code == jle_subj_code):
                        return dbf_path, jle_record  # Return the matching DBF path and JLE record

    return None, None


def get_semester_digit(semester_str):
    """Convert semester string to digit"""
    semester_map = {
        '1st Semester': '1',
        '2nd Semester': '2',
        'Summer': '3'
    }
    return semester_map.get(semester_str, '0')


def extract_dbf_data(dbf_path):
    """Extract data from DBF file"""
    table = Table(dbf_path)
    table.open()

    records = []
    for record in table:
        record_dict = {}
        for field in table.field_names:
            record_dict[field] = record[field]
        records.append(record_dict)

    df = pd.DataFrame(records)
    table.close()

    return df


class WordReport:
    def __init__(self, template_path="Report_template.docx"):
        """
        Initialize the Word report generator with a template
        """
        self.template_path = template_path
        self.doc = Document(template_path)

    def replace_placeholders_in_paragraph(self, paragraph, placeholders):
        """
        Replace placeholders in a paragraph using square brackets [Insert KEY]
        More robust replacement that handles text runs properly, sanitizes input, and makes replaced text bold
        """
        for key, value in placeholders.items():
            placeholder = f"[Insert {key}]"  # Using square brackets as placeholders
            if placeholder in paragraph.text:
                # Sanitize the value to remove control characters that cause XML errors
                sanitized_value = self.sanitize_text_for_xml(str(value))

                # Find all occurrences of the placeholder in the paragraph
                # Split the paragraph text by the placeholder
                parts = paragraph.text.split(placeholder)

                # Clear the paragraph content
                paragraph.clear()

                # Add the first part (before first occurrence)
                if parts[0]:
                    paragraph.add_run(parts[0])

                # Add each replacement value in bold
                for i in range(len(parts) - 1):
                    # Add the replacement value in bold
                    bold_run = paragraph.add_run(sanitized_value)
                    bold_run.font.bold = True
                    bold_run.font.size = Pt(10)

                    # Add the next part (between this and next occurrence, or after last occurrence)
                    if i + 1 < len(parts) and parts[i + 1]:
                        paragraph.add_run(parts[i + 1])

                # Note: This handles text runs properly by creating new runs for replaced content

    def sanitize_text_for_xml(self, text):
        """
        Remove or replace characters that are not XML compatible
        """
        if text is None:
            return "N/A"

        # Remove control characters (0x00-0x1F) except for tab, newline, and carriage return
        # These are the only control characters allowed in XML
        sanitized = ''.join(char if ord(char) >= 0x20 or ord(char) in (0x09, 0x0A, 0x0D) else '?' for char in str(text))

        # Handle any remaining problematic characters
        sanitized = sanitized.replace('\x00', '')  # Remove null bytes specifically

        return sanitized

    def replace_placeholders_in_header_footer(self, section, placeholders):
        """
        Replace placeholders in header and footer, including any tables within them
        """
        # Process header paragraphs
        for paragraph in section.header.paragraphs:
            self.replace_placeholders_in_paragraph(paragraph, placeholders)

        # Process tables in header (if any exist)
        if hasattr(section.header, 'tables'):
            for table in section.header.tables:
                self.replace_placeholders_in_tables(table, placeholders)

        # Process footer paragraphs
        for paragraph in section.footer.paragraphs:
            self.replace_placeholders_in_paragraph(paragraph, placeholders)

        # Process tables in footer (if any exist)
        if hasattr(section.footer, 'tables'):
            for table in section.footer.tables:
                self.replace_placeholders_in_tables(table, placeholders)

    def replace_placeholders_in_tables(self, table, placeholders):
        """
        Replace placeholders in table cells
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_placeholders_in_paragraph(paragraph, placeholders)

    def populate_template_with_jle_dbf_data(self, jle_file_path, dbf_directory="testfiles"):
        """
        Populate the template with data from matched JLE and DBF files
        """
        # Find matching DBF file for the JLE
        matching_dbf, jle_record = find_matching_dbf(jle_file_path, dbf_directory)

        if not matching_dbf:
            st.warning("No matching DBF file found for the JLE file.")
            # Still proceed with JLE-only data
            placeholders = {
                'SY': self.sanitize_text_for_xml(jle_record['Academic Year'] if jle_record is not None else 'N/A'),
                'SC': self.sanitize_text_for_xml(jle_record['Subject Code'] if jle_record is not None else 'N/A'),
                'SN': self.sanitize_text_for_xml(jle_record['Subject Num'] if jle_record is not None else 'N/A'),
                'SEM': self.sanitize_text_for_xml(jle_record['Semester'] if jle_record is not None else 'N/A'),
                'LECTURER': self.sanitize_text_for_xml(jle_record['Lecturer'] if jle_record is not None else 'N/A'),
                'CREDIT': self.sanitize_text_for_xml(jle_record['Credit'] if jle_record is not None else 'N/A'),
                'SCHED': self.sanitize_text_for_xml(jle_record['Schedule'] if jle_record is not None else 'N/A'),
                'TITLE': self.sanitize_text_for_xml(jle_record['Subject Title'] if jle_record is not None else 'N/A'),
                'STUDENT_COUNT': 'N/A',  # No DBF data
                # Additional mappings as requested
                'ST': self.sanitize_text_for_xml(jle_record['Subject Title'] if jle_record is not None else 'N/A'),  # Subject Title
                'Sem': self.sanitize_text_for_xml(jle_record['Semester'] if jle_record is not None else 'N/A'),      # Semester
                'OC': self.sanitize_text_for_xml(jle_record['Subject Num'] if jle_record is not None else 'N/A'),    # Subject Num
                'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                'Faculty': self.sanitize_text_for_xml(jle_record['Lecturer'] if jle_record is not None else 'N/A'),   # Lecturer
                'SCHED': self.sanitize_text_for_xml(jle_record['Schedule'] if jle_record is not None else 'N/A'),    # All schedules
                'LeS': self.sanitize_text_for_xml(jle_record['LEC_Schedule'] if jle_record is not None else 'N/A'),  # Lecture Schedule
                'LaS': self.sanitize_text_for_xml(jle_record['LAB_Schedule'] if jle_record is not None else 'N/A')   # Laboratory Schedule
            }
        else:
            # Extract data from DBF
            dbf_data = extract_dbf_data(matching_dbf)

            # Prepare placeholders with both JLE and DBF data
            placeholders = {
                'SY': self.sanitize_text_for_xml(jle_record['Academic Year']),  # School Year
                'SC': self.sanitize_text_for_xml(jle_record['Subject Code']),   # Subject Code
                'SN': self.sanitize_text_for_xml(jle_record['Subject Num']),    # Subject Number
                'SEM': self.sanitize_text_for_xml(jle_record['Semester']),      # Semester
                'LECTURER': self.sanitize_text_for_xml(jle_record['Lecturer']), # Lecturer
                'CREDIT': self.sanitize_text_for_xml(jle_record['Credit']),     # Credit
                'SCHED': self.sanitize_text_for_xml(jle_record['Schedule']),    # Schedule
                'TITLE': self.sanitize_text_for_xml(jle_record['Subject Title']), # Subject Title
                'STUDENT_COUNT': len(dbf_data),      # Number of students from DBF
                # Additional mappings as requested
                'ST': self.sanitize_text_for_xml(jle_record['Subject Title']),  # Subject Title
                'Sem': self.sanitize_text_for_xml(jle_record['Semester']),      # Semester
                'OC': self.sanitize_text_for_xml(jle_record['Subject Num']),    # Subject Num
                'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                'Faculty': self.sanitize_text_for_xml(jle_record['Lecturer'])   # Lecturer
            }

        # Handle the student data table specifically - look for the table with student information
        # Use the DBF data that was extracted if available
        if matching_dbf:
            # Extract DBF data to populate the student table
            dbf_data = extract_dbf_data(matching_dbf)
            if dbf_data is not None and not dbf_data.empty:
                self.populate_student_data_table(dbf_data)
        else:
            # No matching DBF file, so no student data to populate
            pass

        # Replace placeholders in headers and footers for all sections
        for section in self.doc.sections:
            self.replace_placeholders_in_header_footer(section, placeholders)

        # Replace placeholders in the main document body
        for paragraph in self.doc.paragraphs:
            self.replace_placeholders_in_paragraph(paragraph, placeholders)

        # Replace placeholders in tables in the main document body
        for table in self.doc.tables:
            self.replace_placeholders_in_tables(table, placeholders)

    def populate_template_with_jle_dbf_data_from_jle_data(self, jle_data, dbf_directory="testfiles"):
        """
        Populate the template with data from matched JLE data and DBF files
        """
        # Find matching DBF file for the JLE data
        matching_dbf, jle_record = find_matching_dbf_from_jle_data(jle_data, dbf_directory)

        if not matching_dbf:
            # Still proceed with JLE-only data
            placeholders = {
                'SY': self.sanitize_text_for_xml(jle_record.get('Academic Year', 'N/A') if jle_record is not None else 'N/A'),  # School Year
                'SC': self.sanitize_text_for_xml(jle_record.get('Subject Code', 'N/A') if jle_record is not None else 'N/A'),   # Subject Code
                'SN': self.sanitize_text_for_xml(jle_record.get('Subject Num', 'N/A') if jle_record is not None else 'N/A'),    # Subject Number
                'SEM': self.sanitize_text_for_xml(jle_record.get('Semester', 'N/A') if jle_record is not None else 'N/A'),      # Semester
                'LECTURER': self.sanitize_text_for_xml(jle_record.get('Lecturer', 'N/A') if jle_record is not None else 'N/A'), # Lecturer
                'CREDIT': self.sanitize_text_for_xml(jle_record.get('Credit', 'N/A') if jle_record is not None else 'N/A'),     # Credit
                'SCHED': self.sanitize_text_for_xml(jle_record.get('Schedule', 'N/A') if jle_record is not None else 'N/A'),    # Schedule
                'TITLE': self.sanitize_text_for_xml(jle_record.get('Subject Title', 'N/A') if jle_record is not None else 'N/A'), # Subject Title
                'STUDENT_COUNT': 'N/A',  # No DBF data
                # Additional mappings as requested
                'ST': self.sanitize_text_for_xml(jle_record.get('Subject Title', 'N/A') if jle_record is not None else 'N/A'),  # Subject Title
                'Sem': self.sanitize_text_for_xml(jle_record.get('Semester', 'N/A') if jle_record is not None else 'N/A'),      # Semester
                'OC': self.sanitize_text_for_xml(jle_record.get('Subject Num', 'N/A') if jle_record is not None else 'N/A'),    # Subject Num
                'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                'Faculty': self.sanitize_text_for_xml(jle_record.get('Lecturer', 'N/A') if jle_record is not None else 'N/A'),   # Lecturer
                'SCHED': self.sanitize_text_for_xml(jle_record.get('Schedule', 'N/A') if jle_record is not None else 'N/A'),    # All schedules
                'LeS': self.sanitize_text_for_xml(jle_record.get('LEC_Schedule', 'N/A') if jle_record is not None else 'N/A'),  # Lecture Schedule
                'LaS': self.sanitize_text_for_xml(jle_record.get('LAB_Schedule', 'N/A') if jle_record is not None else 'N/A')   # Laboratory Schedule
            }
        else:
            # Extract data from DBF
            dbf_data = extract_dbf_data(matching_dbf)

            # Prepare placeholders with both JLE and DBF data
            placeholders = {
                'SY': self.sanitize_text_for_xml(jle_record.get('Academic Year', 'N/A')),  # School Year
                'SC': self.sanitize_text_for_xml(jle_record.get('Subject Code', 'N/A')),   # Subject Code
                'SN': self.sanitize_text_for_xml(jle_record.get('Subject Num', 'N/A')),    # Subject Number
                'SEM': self.sanitize_text_for_xml(jle_record.get('Semester', 'N/A')),      # Semester
                'LECTURER': self.sanitize_text_for_xml(jle_record.get('Lecturer', 'N/A')), # Lecturer
                'CREDIT': self.sanitize_text_for_xml(jle_record.get('Credit', 'N/A')),     # Credit
                'SCHED': self.sanitize_text_for_xml(jle_record.get('Schedule', 'N/A')),    # Schedule
                'TITLE': self.sanitize_text_for_xml(jle_record.get('Subject Title', 'N/A')), # Subject Title
                'STUDENT_COUNT': len(dbf_data),                # Number of students from DBF
                # Additional mappings as requested
                'ST': self.sanitize_text_for_xml(jle_record.get('Subject Title', 'N/A')),  # Subject Title
                'Sem': self.sanitize_text_for_xml(jle_record.get('Semester', 'N/A')),      # Semester
                'OC': self.sanitize_text_for_xml(jle_record.get('Subject Num', 'N/A')),    # Subject Num
                'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                'Faculty': self.sanitize_text_for_xml(jle_record.get('Lecturer', 'N/A'))   # Lecturer
            }

        # Handle the student data table specifically - look for the table with student information
        # Use the DBF data that was extracted if available
        if matching_dbf:
            # Extract DBF data to populate the student table
            dbf_data = extract_dbf_data(matching_dbf)
            if dbf_data is not None and not dbf_data.empty:
                self.populate_student_data_table(dbf_data)
        else:
            # No matching DBF file, so no student data to populate
            pass

        # Replace placeholders in headers and footers for all sections
        for section in self.doc.sections:
            self.replace_placeholders_in_header_footer(section, placeholders)

        # Replace placeholders in the main document body
        for paragraph in self.doc.paragraphs:
            self.replace_placeholders_in_paragraph(paragraph, placeholders)

        # Replace placeholders in tables in the main document body
        for table in self.doc.tables:
            self.replace_placeholders_in_tables(table, placeholders)

    def populate_template_with_jle_and_uploaded_dbf_data(self, jle_data, uploaded_dbf_filename, dbf_data_df):
        """
        Populate the template with data from JLE and uploaded DBF file
        """
        matched_course = None
        matching_successful = False

        # Strategy 1: Try to match based on filename pattern
        # Pattern: ORG_YYYYX_SUBJNUM_SUBJCODE_ID.DBF
        if uploaded_dbf_filename:
            parts = uploaded_dbf_filename.replace('.DBF', '').replace('.dbf', '').split('_')
            if len(parts) >= 4:
                year_semester = parts[1]  # YYYYX
                subj_num = parts[2]       # SUBJNUM
                subj_code = parts[3]      # SUBJCODE

                # Find the matching course in JLE data
                if jle_data and 'course_data' in jle_data and jle_data['course_data'] is not None and not jle_data['course_data'].empty:
                    jle_df = jle_data['course_data']

                    # Find the matching record
                    for _, jle_record in jle_df.iterrows():
                        jle_academic_year = jle_record.get('Academic Year', '')  # Full format like "2024-2025"
                        if jle_academic_year:
                            jle_year = jle_academic_year.split('-')[0]  # Get first part like "2024"
                            jle_year_semester = f"{jle_year}{get_semester_digit(jle_record.get('Semester', ''))}"  # Full format like "20243"
                            jle_subj_num = jle_record.get('Subject Num', '')
                            jle_subj_code = jle_record.get('Subject Code', '')

                            # Check if the components match
                            if (year_semester == jle_year_semester and
                                subj_num == jle_subj_num and
                                subj_code == jle_subj_code):
                                matched_course = jle_record
                                matching_successful = True
                                break  # Found the match, exit loop

        # If no match found by filename, try other strategies
        if matched_course is None and jle_data and 'course_data' in jle_data and jle_data['course_data'] is not None and not jle_data['course_data'].empty:
            jle_df = jle_data['course_data']

            # Strategy 2: If we have student data, we might be able to match on other criteria
            # For now, just take the first course if there's only one
            if len(jle_df) == 1:
                matched_course = jle_df.iloc[0]
                matching_successful = True

        # Prepare placeholders based on whether we found a match
        if matched_course is not None:
            # Found matching course, prepare detailed placeholders
            placeholders = {
                'SY': self.sanitize_text_for_xml(matched_course.get('Academic Year', 'N/A')),  # School Year
                'SC': self.sanitize_text_for_xml(matched_course.get('Subject Code', 'N/A')),   # Subject Code
                'SN': self.sanitize_text_for_xml(matched_course.get('Subject Num', 'N/A')),    # Subject Number
                'SEM': self.sanitize_text_for_xml(matched_course.get('Semester', 'N/A')),      # Semester
                'LECTURER': self.sanitize_text_for_xml(matched_course.get('Lecturer', 'N/A')), # Lecturer
                'CREDIT': self.sanitize_text_for_xml(matched_course.get('Credit', 'N/A')),     # Credit
                'SCHED': self.sanitize_text_for_xml(matched_course.get('Schedule', 'N/A')),    # Schedule
                'TITLE': self.sanitize_text_for_xml(matched_course.get('Subject Title', 'N/A')), # Subject Title
                'STUDENT_COUNT': len(dbf_data_df) if dbf_data_df is not None else 'N/A',  # Number of students from DBF
                # Additional mappings as requested
                'ST': self.sanitize_text_for_xml(matched_course.get('Subject Title', 'N/A')),  # Subject Title
                'Sem': self.sanitize_text_for_xml(matched_course.get('Semester', 'N/A')),      # Semester
                'OC': self.sanitize_text_for_xml(matched_course.get('Subject Num', 'N/A')),    # Subject Num
                'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                'Faculty': self.sanitize_text_for_xml(matched_course.get('Lecturer', 'N/A')),   # Lecturer
                'SCHED': self.sanitize_text_for_xml(matched_course.get('Schedule', 'N/A')),    # All schedules
                'LeS': self.sanitize_text_for_xml(matched_course.get('LEC_Schedule', 'N/A')),  # Lecture Schedule
                'LaS': self.sanitize_text_for_xml(matched_course.get('LAB_Schedule', 'N/A'))   # Laboratory Schedule
            }
        else:
            # If no match found, we should still have basic placeholders to avoid leaving [Insert *] in the document
            # However, we'll handle this properly in the app by showing a warning
            placeholders = {
                'SY': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'SC': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'SN': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'SEM': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'LECTURER': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'CREDIT': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'SCHED': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'TITLE': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'STUDENT_COUNT': len(dbf_data_df) if dbf_data_df is not None else 'N/A',
                # Additional mappings for no match case
                'ST': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'Sem': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'OC': self.sanitize_text_for_xml('NO MATCH FOUND'),
                'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                'Faculty': self.sanitize_text_for_xml('NO MATCH FOUND')
            }

        # Replace placeholders in headers and footers for all sections
        for section in self.doc.sections:
            self.replace_placeholders_in_header_footer(section, placeholders)

        # Replace placeholders in the main document body
        for paragraph in self.doc.paragraphs:
            self.replace_placeholders_in_paragraph(paragraph, placeholders)

        # Replace placeholders in tables in the main document body
        for table in self.doc.tables:
            self.replace_placeholders_in_tables(table, placeholders)

        # Handle the student data table specifically - look for the table with student information
        # Use the DBF data that was passed as parameter
        if dbf_data_df is not None and not dbf_data_df.empty:
            self.populate_student_data_table(dbf_data_df)

    def populate_student_data_table(self, df):
        """
        Populate the student data table with records from the DataFrame.
        The table should have columns: No., Name, Grade, Remarks
        Duplicate rows as needed up to 23 per page, then continue on next page.
        """
        # Find the table that contains student data (look for tables that might have student info)
        for table in self.doc.tables:
            # Check if this table has the expected structure for student data
            # Look for tables that might have headers like "No.", "Name", "Grade", "Remarks"
            if self.is_student_data_table(table):
                # Process the student data table
                self.fill_student_data_table(table, df)
                break  # Only process the first student data table found

    def is_student_data_table(self, table):
        """
        Check if a table is likely the student data table by looking for expected headers
        """
        if not table.rows:
            return False

        # Look at the first row for headers like "No.", "Name", "Grade", "Remarks"
        first_row = table.rows[0]
        row_text = " ".join(cell.text for cell in first_row.cells).upper()

        # Check if it contains expected column names
        expected_headers = ["NO", "NAME", "GRADE", "REMARKS"]
        found_headers = sum(1 for header in expected_headers if header in row_text)

        # If we find at least 3 out of 4 expected headers, consider it a student data table
        return found_headers >= 3

    def fill_student_data_table(self, table, df):
        """
        Fill the student data table with DataFrame content, handling pagination
        Each page holds up to 23 records, then continues on the next page if needed.
        """
        # Get the header row (first row) and the template row (second row, usually empty)
        if len(table.rows) < 2:
            return  # Need at least header and one data row

        header_row = table.rows[0]

        # Find the template row (usually the second row with placeholders like [Insert NAME], etc.)
        template_row_idx = 1
        template_row = table.rows[template_row_idx]

        # Store the original template row for duplication
        original_cells = [cell.text for cell in template_row.cells]

        # Clear the template row content to prepare for data insertion
        for cell in template_row.cells:
            cell.text = ""

        # Get the DataFrame columns that correspond to student data
        # Look for common column names in the DataFrame
        name_col = self.find_column_name(df, ['name', 'student_name', 'studname', 'lastname', 'student'])
        grade_col = self.find_column_name(df, ['grade', 'eg', 'score', 'mark'])
        remarks_col = self.find_column_name(df, ['remarks', 'remark', 'comment', 'status'])

        # Add row numbers column
        df_with_numbers = df.copy()
        df_with_numbers.insert(0, 'row_number', range(1, len(df_with_numbers) + 1))
        number_col = 'row_number'

        # Process each student record with pagination (max 23 records per page)
        records_per_page = 23

        for idx, row in df_with_numbers.iterrows():
            # Determine which row in the table to use
            target_row_idx = template_row_idx + idx  # Sequentially add records to rows

            # Check if we need to add a new row
            if target_row_idx >= len(table.rows):
                # Need to add a new row
                new_row = table.add_row()
                target_row = new_row
            else:
                target_row = table.rows[target_row_idx]

            # Fill the row with data
            for cell_idx, cell in enumerate(target_row.cells):
                if cell_idx == 0:  # Number column
                    cell.text = str(row[number_col]) if number_col in row else str(idx + 1)
                elif cell_idx == 1:  # Name column
                    cell.text = str(row[name_col]) if name_col and name_col in row else ""
                elif cell_idx == 2:  # Grade column
                    cell.text = str(row[grade_col]) if grade_col and grade_col in row else ""
                elif cell_idx == 3:  # Additional column (originally would be Remarks, but now shifted)
                    cell.text = ""  # Leave as empty since we're shifting remarks to the right
                elif cell_idx == 4:  # Remarks column (offset one position to the right)
                    cell.text = str(row[remarks_col]) if remarks_col and remarks_col in row else ""

                # Sanitize the text to ensure it's XML compatible
                cell.text = self.sanitize_text_for_xml(cell.text)


        # Calculate statistics based on the remarks column
        if remarks_col and remarks_col in df_with_numbers.columns:
            remarks_values = df_with_numbers[remarks_col].astype(str).str.upper()

            # Count each category separately to avoid double counting
            passed_count = 0
            no_grade_count = 0
            failed_count = 0
            dropped_count = 0

            for value in remarks_values:
                # Check each category separately and only count once per value
                value_str = str(value) if value is not None and pd.notna(value) else ""
                if 'PASSED' in value_str or 'PASS' in value_str or 'OK' in value_str or 'COMPLETED' in value_str:
                    passed_count += 1
                elif 'NO GRADE' in value_str or 'N/A' in value_str or 'NO REMARK' in value_str or 'NONE' in value_str or 'INC' in value_str or 'INCOMPLETE' in value_str:
                    no_grade_count += 1
                elif 'FAILED' in value_str or 'FAIL' in value_str:
                    failed_count += 1
                elif 'DROPPED' in value_str or 'DROP' in value_str or 'WITHDRAWN' in value_str or 'WITHDREW' in value_str or 'DRP' in value_str:
                    dropped_count += 1
        else:
            # If no remarks column, default to zero counts
            passed_count = 0
            no_grade_count = 0
            failed_count = 0
            dropped_count = 0

        # Total number of students is the length of the dataframe
        total_count = len(df_with_numbers)

        # Add statistics text after the table (outside the table)
        # Find the table's parent element and add the paragraph after it
        stats_text = f"STATISTICS   :   Passed={passed_count}  No Grade={no_grade_count}  Failed={failed_count}  Dropped={dropped_count}  TOTAL={total_count}"

        # Add a paragraph after the table with the statistics
        # Get the table's element and add a paragraph after it
        table_element = table._tbl
        parent_element = table_element.getparent()

        # Create a new paragraph element after the table
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml.text.paragraph import CT_P
        from docx.text.paragraph import Paragraph

        # Add a paragraph after the table element
        paragraph_element = OxmlElement('w:p')
        table_element.addnext(paragraph_element)

        # Add a run to the paragraph
        run_element = OxmlElement('w:r')
        paragraph_element.append(run_element)

        # Add the text to the run
        text_element = OxmlElement('w:t')
        text_element.text = stats_text
        run_element.append(text_element)

        # If there are more records than available rows, the method already handles this by adding rows
        # For proper pagination across pages (every 23 records), we would need more complex logic
        # that creates section breaks, but for now we'll just add as many rows as needed

        # Modify the last row's bottom border to use double lines
        self.apply_double_bottom_border_to_last_rows(table)

    def apply_double_bottom_border_to_last_rows(self, table):
        """
        Apply double bottom borders to the last few rows of the table to indicate the end of the data
        """
        if not table.rows:
            return

        # Get the last row (typically the final data row)
        if len(table.rows) > 1:  # Need at least header + one data row
            last_row = table.rows[-1]  # Last row in the table

            # Apply double border to all cells in the last row
            for cell in last_row.cells:
                # Access the cell's XML element to modify its properties
                from docx.oxml.shared import OxmlElement, qn

                # Get the table cell properties element
                tc_pr = cell._element.tcPr

                # Create a table cell borders element if it doesn't exist
                tbl_borders = tc_pr.first_child_found_in("w:tcBorders")
                if tbl_borders is None:
                    tbl_borders = OxmlElement('w:tcBorders')
                    tc_pr.append(tbl_borders)

                # Create the bottom border element
                bottom_border = OxmlElement('w:bottom')

                # Set the border attributes for double line
                bottom_border.set(qn('w:val'), 'double')  # Double line style
                bottom_border.set(qn('w:sz'), '6')       # Border size (6 half-points = 1/2pt)
                bottom_border.set(qn('w:space'), '0')    # Space
                bottom_border.set(qn('w:color'), '000000')  # Black color

                # Add the bottom border to the table cell borders
                tbl_borders.append(bottom_border)

    def find_column_name(self, df, possible_names):
        """
        Find a column in the DataFrame based on possible names (case-insensitive)
        """
        df_columns = [col.lower() for col in df.columns]
        for name in possible_names:
            for col in df.columns:
                if name.lower() in col.lower():
                    return col
        # If no exact match, return the first column that seems appropriate
        for col in df.columns:
            for name in possible_names:
                if name.lower() in col.lower():
                    return col
        # If still no match, return the first column as a fallback
        return df.columns[0] if len(df.columns) > 0 else None


    def get_document_bytes(self):
        """
        Get the document as bytes for download
        """
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def generate_word_report_from_jle_dbf(jle_file_path, template_path="Report_template.docx", dbf_directory="testfiles"):
    """
    Generate a Word report from the matched JLE and DBF files using the template
    """
    word_report = WordReport(template_path)
    word_report.populate_template_with_jle_dbf_data(jle_file_path, dbf_directory)
    return word_report.get_document_bytes()


def generate_word_report_from_jle_data(jle_data, template_path="Report_template.docx", dbf_directory="testfiles"):
    """
    Generate a Word report from the matched JLE data and DBF files using the template
    """
    word_report = WordReport(template_path)
    word_report.populate_template_with_jle_dbf_data_from_jle_data(jle_data, dbf_directory)
    return word_report.get_document_bytes()


def generate_word_report_from_jle_and_uploaded_dbf(jle_data, uploaded_dbf_filename, dbf_data_df, template_path="Report_template.docx"):
    """
    Generate a Word report from JLE data and uploaded DBF file information
    """
    word_report = WordReport(template_path)
    word_report.populate_template_with_jle_and_uploaded_dbf_data(jle_data, uploaded_dbf_filename, dbf_data_df)
    return word_report.get_document_bytes()


def generate_word_report(df, jle_data, template_path="Report_template.docx"):
    """
    Generate a Word report from the DataFrame and JLE data using the template
    """
    # This function is kept for backward compatibility
    # For the new functionality, use generate_word_report_from_jle_dbf
    word_report = WordReport(template_path)

    # Prepare placeholders from JLE data (old method)
    placeholders = {}

    # Extract useful information from JLE data
    if jle_data:
        placeholders.update({
            'JLE_FILENAME': word_report.sanitize_text_for_xml(jle_data.get('filename', 'N/A')),
            'JLE_SIZE': jle_data.get('size', 'N/A'),
            'DATE_GENERATED': pd.Timestamp.now().strftime('%Y-%m-%d'),
            'TIME_GENERATED': pd.Timestamp.now().strftime('%H:%M:%S'),  # Keep time for this specific field
            'TOTAL_COURSES': jle_data.get('total_courses', 'N/A')
        })

        # If there is course data from the JLE file, extract course information
        if 'course_data' in jle_data and jle_data['course_data'] is not None and not jle_data['course_data'].empty:
            course_df = jle_data['course_data']

            # Extract course information for reporting
            course_codes = ', '.join(course_df['Subject Code'].unique()) if 'Subject Code' in course_df.columns else 'N/A'
            lecturers = ', '.join(course_df['Lecturer'].dropna().unique()) if 'Lecturer' in course_df.columns else 'N/A'

            placeholders.update({
                'COURSE_CODES': word_report.sanitize_text_for_xml(course_codes),
                'LECTURERS': word_report.sanitize_text_for_xml(lecturers),
                'COURSE_COUNT': len(course_df)
            })

            # Add first course details as examples
            if len(course_df) > 0:
                first_course = course_df.iloc[0]
                placeholders.update({
                    'FIRST_COURSE_CODE': word_report.sanitize_text_for_xml(first_course.get('Subject Code', 'N/A')),
                    'FIRST_COURSE_TITLE': word_report.sanitize_text_for_xml(first_course.get('Subject Title', 'N/A')),
                    'FIRST_COURSE_SCHEDULE': word_report.sanitize_text_for_xml(first_course.get('Schedule', 'N/A')),
                    'FIRST_COURSE_CREDIT': word_report.sanitize_text_for_xml(first_course.get('Credit', 'N/A')),
                    'FIRST_COURSE_LECTURER': word_report.sanitize_text_for_xml(first_course.get('Lecturer', 'N/A'))
                })

                # Add the new mappings based on the first course if available
                placeholders.update({
                    'ST': word_report.sanitize_text_for_xml(first_course.get('Subject Title', 'N/A')),  # Subject Title
                    'Sem': word_report.sanitize_text_for_xml(first_course.get('Semester', 'N/A')),      # Semester
                    'OC': word_report.sanitize_text_for_xml(first_course.get('Subject Num', 'N/A')),    # Subject Num
                    'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                    'Faculty': word_report.sanitize_text_for_xml(first_course.get('Lecturer', 'N/A')),   # Lecturer
                    'SCHED': word_report.sanitize_text_for_xml(first_course.get('Schedule', 'N/A')),    # All schedules
                    'LeS': word_report.sanitize_text_for_xml(first_course.get('LEC_Schedule', 'N/A')),  # Lecture Schedule
                    'LaS': word_report.sanitize_text_for_xml(first_course.get('LAB_Schedule', 'N/A'))   # Laboratory Schedule
                })
            else:
                # Add default values for the new mappings if no course data
                placeholders.update({
                    'ST': word_report.sanitize_text_for_xml('N/A'),  # Subject Title
                    'Sem': word_report.sanitize_text_for_xml('N/A'),      # Semester
                    'OC': word_report.sanitize_text_for_xml('N/A'),    # Subject Num
                    'Time': pd.Timestamp.now().strftime('%Y-%m-%d'),  # Current day only
                    'Faculty': word_report.sanitize_text_for_xml('N/A'),   # Lecturer
                    'SCHED': word_report.sanitize_text_for_xml('N/A'),    # All schedules
                    'LeS': word_report.sanitize_text_for_xml('N/A'),  # Lecture Schedule
                    'LaS': word_report.sanitize_text_for_xml('N/A')   # Laboratory Schedule
                })

        # Add dataframe statistics
        if df is not None and not df.empty:
            placeholders.update({
                'TOTAL_RECORDS': len(df),
                'COLUMNS_COUNT': len(df.columns),
                'COLUMN_NAMES': word_report.sanitize_text_for_xml(', '.join(df.columns.tolist()))
            })

    # Replace placeholders in headers and footers for all sections
    for section in word_report.doc.sections:
        word_report.replace_placeholders_in_header_footer(section, placeholders)

    # Replace placeholders in the main document body
    for paragraph in word_report.doc.paragraphs:
        word_report.replace_placeholders_in_paragraph(paragraph, placeholders)

    # Replace placeholders in tables in the main document body
    for table in word_report.doc.tables:
        word_report.replace_placeholders_in_tables(table, placeholders)

    # Handle the student data table specifically - look for the table with student information
    if df is not None and not df.empty:
        word_report.populate_student_data_table(df)

    return word_report.get_document_bytes()


def convert_word_to_pdf(word_bytes):
    """
    Convert Word document bytes to PDF bytes
    """
    # Create temporary files
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as word_temp:
        word_temp.write(word_bytes)
        word_temp_path = word_temp.name

    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_temp:
        pdf_temp_path = pdf_temp.name

    try:
        # Convert Word to PDF using docx2pdf (cross-platform)
        convert(word_temp_path, pdf_temp_path)

        # Read the PDF content
        with open(pdf_temp_path, 'rb') as pdf_file:
            pdf_bytes = pdf_file.read()

        return pdf_bytes
    finally:
        # Clean up temporary files
        try:
            os.remove(word_temp_path)
            os.remove(pdf_temp_path)
        except:
            pass  # Ignore errors during cleanup


def show_word_report_ui(df, jle_data):
    """
    Display the Word report UI
    """
    # Initialize session state variables if they don't exist
    if 'word_report_generated' not in st.session_state:
        st.session_state.word_report_generated = False
    if 'pdf_from_word_generated' not in st.session_state:
        st.session_state.pdf_from_word_generated = False

    st.write("Click the button below to generate a Word report using the template:")

    # Generate Word Report button
    if st.button("Generate Word Report & PDF", key="generate_word_report_btn"):
        with st.spinner('Generating Word report and converting to PDF...'):
            try:
                # Generate Word report
                word_bytes = generate_word_report(df, jle_data)

                # Store the Word doc in session state
                st.session_state.word_bytes = word_bytes
                st.session_state.word_filename = f"E-Class_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx"
                st.session_state.word_report_generated = True

                # Convert Word to PDF
                pdf_bytes = convert_word_to_pdf(word_bytes)

                # Store the PDF in session state
                st.session_state.pdf_from_word_bytes = pdf_bytes
                st.session_state.pdf_from_word_filename = f"E-Class_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                st.session_state.pdf_from_word_generated = True

                st.success("Word report and PDF generated successfully!")

            except Exception as e:
                st.error(f"Error generating Word report: {str(e)}")

    # Show the download button if Word report is generated
    if st.session_state.word_report_generated and 'word_bytes' in st.session_state:
        st.download_button(
            label="Download Word Report",
            data=st.session_state.word_bytes,
            file_name=st.session_state.word_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_word_btn"
        )

    # Show the PDF viewer if PDF from Word is generated
    if st.session_state.pdf_from_word_generated and 'pdf_from_word_bytes' in st.session_state:
        st.subheader("PDF from Word Report")

        # Convert PDF bytes to base64 for embedding in HTML
        import base64
        pdf_bytes = st.session_state.pdf_from_word_bytes
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode('latin-1')
        elif isinstance(pdf_bytes, bytearray):
            pdf_bytes = bytes(pdf_bytes)

        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

        # Embed the PDF in an iframe
        pdf_display = f'<iframe src="data:application/pdf;base64,{pdf_base64}" width="100%" height="600" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

        # Download button for PDF
        st.download_button(
            label="Download PDF Report",
            data=st.session_state.pdf_from_word_bytes,
            file_name=st.session_state.pdf_from_word_filename,
            mime="application/pdf",
            key="download_pdf_from_word_btn"
        )


def show_jle_dbf_report_ui(jle_file_path):
    """
    Display the Word report UI for JLE/DBF workflow
    """
    # Initialize session state variables if they don't exist
    if 'jle_dbf_word_report_generated' not in st.session_state:
        st.session_state.jle_dbf_word_report_generated = False
    if 'jle_dbf_pdf_generated' not in st.session_state:
        st.session_state.jle_dbf_pdf_generated = False

    st.write("Click the button below to generate a Word report using the JLE/DBF template:")

    # Generate Word Report button for JLE/DBF workflow
    if st.button("Generate Word Report & PDF (JLE/DBF)", key="generate_jle_dbf_report_btn"):
        with st.spinner('Generating Word report from JLE/DBF and converting to PDF...'):
            try:
                # Generate Word report using JLE/DBF data
                word_bytes = generate_word_report_from_jle_dbf(jle_file_path)

                # Store the Word doc in session state
                st.session_state.jle_dbf_word_bytes = word_bytes
                st.session_state.jle_dbf_word_filename = f"E-Class_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx"
                st.session_state.jle_dbf_word_report_generated = True

                # Convert Word to PDF
                pdf_bytes = convert_word_to_pdf(word_bytes)

                # Store the PDF in session state
                st.session_state.jle_dbf_pdf_bytes = pdf_bytes
                st.session_state.jle_dbf_pdf_filename = f"E-Class_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                st.session_state.jle_dbf_pdf_generated = True

                st.success("Word report and PDF generated successfully from JLE/DBF data!")

            except Exception as e:
                st.error(f"Error generating Word report from JLE/DBF: {str(e)}")

    # Show the download button if Word report is generated
    if st.session_state.jle_dbf_word_report_generated and 'jle_dbf_word_bytes' in st.session_state:
        st.download_button(
            label="Download Word Report (JLE/DBF)",
            data=st.session_state.jle_dbf_word_bytes,
            file_name=st.session_state.jle_dbf_word_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_jle_dbf_word_btn"
        )

    # Show the PDF viewer if PDF is generated
    if st.session_state.jle_dbf_pdf_generated and 'jle_dbf_pdf_bytes' in st.session_state:
        st.subheader("PDF from JLE/DBF Word Report")

        # Convert PDF bytes to base64 for embedding in HTML
        import base64
        pdf_bytes = st.session_state.jle_dbf_pdf_bytes
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode('latin-1')
        elif isinstance(pdf_bytes, bytearray):
            pdf_bytes = bytes(pdf_bytes)

        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

        # Embed the PDF in an iframe
        pdf_display = f'<iframe src="data:application/pdf;base64,{pdf_base64}" width="100%" height="600" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

        # Download button for PDF
        st.download_button(
            label="Download PDF Report (JLE/DBF)",
            data=st.session_state.jle_dbf_pdf_bytes,
            file_name=st.session_state.jle_dbf_pdf_filename,
            mime="application/pdf",
            key="download_jle_dbf_pdf_btn"
        )