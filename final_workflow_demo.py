#!/usr/bin/env python3
"""
Final demonstration of the complete workflow:
1. Match JLE and DBF files based on naming patterns
2. Extract data from both files
3. Replace placeholders in DOCX template
4. Generate output without modifying the original template
"""
import os
import re
import pandas as pd
from dbf import Table
import shutil
from docx import Document
from docx.shared import Inches
import tempfile


def parse_jle_with_filename_fixed(file_path):
    """
    Fixed version of the JLE parser that properly handles lecturer and credit extraction
    """
    # 1. Extract Semester/Year from Filename
    filename = os.path.basename(file_path)

    # Logic: Look for the pattern "20243" inside "DSO_20243_565.JLE"
    # This assumes the digit immediately following the year is the semester.
    term_map = {'1': '1st Semester', '2': '2nd Semester', '3': 'Summer'}

    # Default values
    academic_year = "Unknown"
    semester = "Unknown"

    # Extract digits (e.g., finding "20243")
    meta_match = re.search(r"(\d{4})(\d)", filename)
    if meta_match:
        year_part = meta_match.group(1)  # 2024
        term_part = meta_match.group(2)  # 3

        academic_year = f"{year_part}-{int(year_part)+1}"
        semester = term_map.get(term_part, f"{term_part}th Term")

    # 2. Extract data from the JLE file content
    with open(file_path, 'rb') as f:
        jle_bytes = f.read()

    # Read the JLE file content as binary and decode with latin1 encoding
    raw_data = jle_bytes.decode('latin1', errors='ignore')

    # Define the "Start of Record" pattern
    record_start_pattern = re.compile(r"(\d{4})\s+([A-Z][A-Z0-9]{2,10})")

    # Find all iterators for the start of records
    matches = list(record_start_pattern.finditer(raw_data))

    parsed_records = []

    for i in range(len(matches)):
        # Determine the start and end of the current record's text chunk
        start_idx = matches[i].start()

        # If there is a next match, end before it starts; otherwise go to EOF
        if i + 1 < len(matches):
            end_idx = matches[i+1].start()
        else:
            end_idx = len(raw_data)

        # Extract the raw chunk for this specific subject
        chunk = raw_data[start_idx:end_idx]

        # --- Extraction Logic ---

        # 1. Extract Subject Num and Code from the match itself
        subj_num = matches[i].group(1)
        full_subj_code = matches[i].group(2)

        # The first letter of the subject code is actually part of the subject number
        subj_num_with_letter = subj_num + full_subj_code[0]  # 2506 + F = 2506F
        real_subj_code = full_subj_code[1:]  # BACC104 (everything after the first letter)

        # Remove the header (Num + Code) from the chunk to process the rest
        rest_of_text = chunk[len(matches[i].group(0)):].replace('\n', ' ').strip()

        # 2. Extract Credit and Lecturer using improved pattern matching
        # The lecturer and credit appear at the end of the record
        # Updated pattern to handle control characters and special formatting
        lecturer_pattern = re.search(r"(\d)\s+([A-Z][A-Z\s\.\-]+?)(?:\x00|\x01|\x02|\x03|\x04|\x05|\x06|\x07|\x08|\x09|\x0a|\x0b|\x0c|\x0d|\x0e|\x0f|\x10|\x11|\x12|\x13|\x14|\x15|\x16|\x17|\x18|\x19|\x1a|\x1b|\x1c|\x1d|\x1e|\x1f|\x7f|$|[^A-Z\s\.\-])", rest_of_text)
        
        credit = None
        lecturer = None

        if lecturer_pattern:
            credit = lecturer_pattern.group(1)
            lecturer = lecturer_pattern.group(2).strip()
        else:
            # Try another approach: look for credit (digit) followed by lecturer name at the end
            # Pattern: digit followed by spaces and then uppercase letters (lecturer name)
            alt_pattern = re.search(r"(\d)\s+([A-Z][A-Z\s\.\-]*[A-Z])(?:\s*$|[\x00-\x1f\x7f]+.*$)", rest_of_text)
            if alt_pattern:
                credit = alt_pattern.group(1)
                lecturer = alt_pattern.group(2).strip()

        # 3. Extract Schedule(s)
        # Pattern: Time (e.g., 730AM-1200PM), Day (e.g., SuSa), Room (e.g., B63)
        sched_pattern_regex = r"(\d{1,4}[AP]M-\s*\d{1,4}[AP]M\s+[A-Za-z]+\s+[A-Z0-9]+)"
        schedules_found = re.findall(sched_pattern_regex, rest_of_text)

        # Clean up the schedule string for the final output
        schedule_str = " / ".join([s.strip() for s in schedules_found])

        # 4. Extract Subject Title
        # The title is whatever is left after removing the schedules and lecturer info
        title_clean = re.sub(sched_pattern_regex, '', rest_of_text)
        # Remove lecturer and credit info from title
        if credit and lecturer:
            # Remove the credit and lecturer from the text
            lecturer_part = f"{credit}\\s+{lecturer}"
            title_clean = re.sub(lecturer_part, '', title_clean, flags=re.IGNORECASE)
        title_clean = " ".join(title_clean.split()) # Remove extra whitespace

        # Clean up title to remove any control characters
        title_clean = re.sub(r'[\x00-\x1f\x7f]', ' ', title_clean).strip()
        if lecturer:
            # Remove lecturer name from title if it appears there
            title_clean = re.sub(re.escape(lecturer), '', title_clean, flags=re.IGNORECASE).strip()

        parsed_records.append({
            "Subject Num": subj_num_with_letter,  # Now includes the letter (e.g., 2506F)
            "Subject Code": real_subj_code,       # The actual subject code (e.g., BACC104)
            "Subject Title": title_clean,
            "Schedule": schedule_str,
            "Credit": credit,
            "Lecturer": lecturer,
            "Academic Year": academic_year,  # Added from filename
            "Semester": semester             # Added from filename
        })

    # Create a DataFrame from the parsed records
    jle_df = pd.DataFrame(parsed_records) if parsed_records else pd.DataFrame()

    return jle_df


def extract_jle_metadata(jle_file_path):
    """Extract metadata from JLE file using the fixed parser"""
    df = parse_jle_with_filename_fixed(jle_file_path)  # Use the fixed parser
    
    if df.empty:
        return None
    
    # Get unique values from the JLE file
    academic_year = df['Academic Year'].iloc[0] if 'Academic Year' in df.columns else "Unknown"
    semester = df['Semester'].iloc[0] if 'Semester' in df.columns else "Unknown"
    
    # Create a mapping of subject info for matching
    subject_info = {}
    for _, row in df.iterrows():
        key = f"{row.get('Academic Year', '')}_{row.get('Subject Num', '')}_{row.get('Subject Code', '')}"
        subject_info[key] = {
            'Subject Number': row.get('Subject Num', ''),
            'Subject Code': row.get('Subject Code', ''),
            'Subject Title': row.get('Subject Title', ''),
            'Schedule': row.get('Schedule', ''),
            'Credit': row.get('Credit', ''),
            'Lecturer': row.get('Lecturer', ''),
            'Academic Year': row.get('Academic Year', ''),
            'Semester': row.get('Semester', '')
        }
    
    return {
        'academic_year': academic_year,
        'semester': semester,
        'subject_info': subject_info,
        'all_data': df
    }


def find_matching_dbf(jle_metadata, dbf_directory="testfiles"):
    """Find DBF file that matches the JLE metadata"""
    if not jle_metadata:
        return None
    
    dbf_files = [f for f in os.listdir(dbf_directory) if f.lower().endswith('.dbf')]
    
    for dbf_file in dbf_files:
        dbf_path = os.path.join(dbf_directory, dbf_file)
        
        # Extract parts from DBF filename to match against JLE
        # Pattern: ORG_YYYYX_SUBJNUM_SUBJCODE_ID.DBF
        parts = dbf_file.replace('.DBF', '').split('_')
        if len(parts) >= 4:
            year_semester = parts[1]  # YYYYX
            subj_num = parts[2]       # SUBJNUM
            subj_code = parts[3]      # SUBJCODE
            
            print(f"Checking DBF: {dbf_file} -> Year/Sem: {year_semester}, SubjNum: {subj_num}, SubjCode: {subj_code}")
            
            # Check if any of the JLE subject info matches
            for jle_key, jle_data in jle_metadata['subject_info'].items():
                jle_academic_year = jle_data['Academic Year']  # Full format like "2024-2025"
                jle_year = jle_academic_year.split('-')[0]  # Get first part like "2024"
                jle_year_semester = f"{jle_year}{get_semester_digit(jle_data['Semester'])}"  # Full format like "20243"
                jle_subj_num = jle_data['Subject Number']
                jle_subj_code = jle_data['Subject Code']
                
                print(f"  Comparing with JLE: Year/Sem: {jle_year_semester}, SubjNum: {jle_subj_num}, SubjCode: {jle_subj_code}")
                
                # Check if the components match
                if (year_semester == jle_year_semester and 
                    subj_num == jle_subj_num and 
                    subj_code == jle_subj_code):
                    print(f"  MATCH FOUND!")
                    return dbf_path, jle_data
    
    return None, None


def get_semester_digit(semester_str):
    """Convert semester string to digit"""
    semester_map = {
        '1st Semester': '1',
        '2nd Semester': '2',
        'Summer': '3'
    }
    return semester_map.get(semester_str, '0')


def extract_dbf_data(dbf_path):
    """Extract data from DBF file"""
    table = Table(dbf_path)
    table.open()
    
    records = []
    for record in table:
        record_dict = {}
        for field in table.field_names:
            record_dict[field] = record[field]
        records.append(record_dict)
    
    df = pd.DataFrame(records)
    table.close()
    
    return df


def replace_placeholders(template_path, output_path, replacement_data):
    """Replace placeholders in DOCX template without modifying the original"""
    # Load the template
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacement_data.items():
            placeholder = f"[Insert {key}]"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacement_data.items():
                    placeholder = f"[Insert {key}]"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    # Save the modified document to the output path
    doc.save(output_path)


def demonstrate_complete_workflow():
    """Demonstrate the complete workflow"""
    print("COMPLETE WORKFLOW DEMONSTRATION")
    print("=" * 60)
    
    # Use the test files
    jle_file_path = "testfiles/DSO_20243_565.JLE"
    template_path = "Report_template.docx"
    
    # Check if files exist
    if not os.path.exists(jle_file_path):
        print(f"ERROR: JLE file {jle_file_path} does not exist!")
        return
    
    if not os.path.exists(template_path):
        print(f"ERROR: Template file {template_path} does not exist!")
        return
    
    print(f"Step 1: Processing JLE file: {jle_file_path}")
    
    # Step 1: Extract metadata from JLE
    jle_metadata = extract_jle_metadata(jle_file_path)
    if not jle_metadata:
        print("ERROR: Could not extract metadata from JLE file")
        return
    
    print(f"  - Academic Year: {jle_metadata['academic_year']}")
    print(f"  - Semester: {jle_metadata['semester']}")
    print(f"  - Found {len(jle_metadata['subject_info'])} subjects in JLE")
    
    # Step 2: Find matching DBF file
    print(f"\nStep 2: Finding matching DBF file...")
    matching_dbf, jle_subject_data = find_matching_dbf(jle_metadata)
    
    if matching_dbf:
        print(f"  - Found matching DBF file: {matching_dbf}")
        
        # Step 3: Extract data from DBF
        print(f"\nStep 3: Extracting data from DBF file...")
        dbf_data = extract_dbf_data(matching_dbf)
        print(f"  - DBF records count: {len(dbf_data)}")
        print(f"  - DBF columns: {list(dbf_data.columns)}")
        
        # Step 4: Prepare replacement data
        print(f"\nStep 4: Preparing replacement data...")
        replacement_data = {
            'SY': jle_subject_data['Academic Year'],  # School Year
            'SC': jle_subject_data['Subject Code'],   # Subject Code
            'SN': jle_subject_data['Subject Number'], # Subject Number
            'SEM': jle_subject_data['Semester'],      # Semester
            'LECTURER': jle_subject_data['Lecturer'], # Lecturer
            'CREDIT': jle_subject_data['Credit'],     # Credit
            'SCHED': jle_subject_data['Schedule'],    # Schedule
            'TITLE': jle_subject_data['Subject Title'], # Subject Title
            'STUDENT_COUNT': len(dbf_data)            # Number of students
        }
        
        print(f"  Replacement data prepared:")
        for key, value in replacement_data.items():
            print(f"    [Insert {key}]: {value}")
        
        # Step 5: Replace placeholders in template (without modifying original)
        print(f"\nStep 5: Replacing placeholders in template...")
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_path = temp_file.name
        
        try:
            replace_placeholders(template_path, temp_path, replacement_data)
            print(f"  - Template processed successfully!")
            print(f"  - Output saved to: {temp_path}")
            print(f"  - Original template ({template_path}) was NOT modified")
            
            # Show the expected replacements
            print(f"\nEXPECTED REPLACEMENTS:")
            print(f"  [Insert SY] will be replaced with: {replacement_data['SY']}")
            print(f"  [Insert SC] will be replaced with: {replacement_data['SC']}")
            print(f"  [Insert SN] will be replaced with: {replacement_data['SN']}")
            print(f"  [Insert SEM] will be replaced with: {replacement_data['SEM']}")
            print(f"  [Insert LECTURER] will be replaced with: {replacement_data['LECTURER']}")
            print(f"  [Insert CREDIT] will be replaced with: {replacement_data['CREDIT']}")
            print(f"  [Insert SCHED] will be replaced with: {replacement_data['SCHED']}")
            print(f"  [Insert TITLE] will be replaced with: {replacement_data['TITLE']}")
            print(f"  [Insert STUDENT_COUNT] will be replaced with: {replacement_data['STUDENT_COUNT']}")
            
            print(f"\nWORKFLOW COMPLETE: Template is ready for PDF conversion!")
            
        except Exception as e:
            print(f"ERROR processing template: {e}")
            import traceback
            traceback.print_exc()
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    else:
        print(f"\nNo matching DBF file found for the JLE data")
        print("Available JLE subject keys:")
        for key in jle_metadata['subject_info'].keys():
            print(f"  - {key}")


if __name__ == "__main__":
    demonstrate_complete_workflow()